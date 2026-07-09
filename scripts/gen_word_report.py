#!/usr/bin/env python3
"""
Generate Word reports:
  1. Full report (all statuses)
  2. Active-only report (故障 / 維修中 / 待確認)
Both with complete English translations.
"""
import urllib.request, json, datetime, sys, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ───────────────────────────────────────────────────────────────────
API_KEY    = "AIzaSyBRNMZFCSnWk1X_HZMuDa_ym-Zvwk9ei-U"
PROJECT_ID = "fics-6e2cd"
BASE_URL   = f"https://firestore.googleapis.com/v1/projects/{PROJECT_ID}/databases/(default)/documents"
TODAY      = datetime.date.today().isoformat()
# 輸出目錄：預設為腳本所在目錄，可用環境變數 REPORT_OUT_DIR 覆寫
SCRATCHPAD = os.environ.get("REPORT_OUT_DIR", os.path.dirname(os.path.abspath(__file__)))

# ── Rotations (57 cars) ───────────────────────────────────────────────────────
ROTATIONS = [
    ("第一輪 1st Rotation",  ["CN360","CN359","CN358","CN357","CN356","CN353","CN370","CN369","NMS382","NMS381","NMS383"]),
    ("第二輪 2nd Rotation",  ["CN362","CN320","CN321","CN309","CN302","CN305","CN301","CN344","CN329","CN323","CN303","CN319","CN308","CN310","CN314"]),
    ("第三輪 3rd Rotation",  ["CN324","CN337","CN327","CN347","CN326","CN346","CN361","CN355","CN342","CN311","CN330","CN341","CN348","CN354","CN331"]),
    ("第四輪 4th Rotation",  ["CN343","CN345","CN312","CN339","CN340","CN336","CN307","CN313","CN364","CN322","CN325","CN306","CN338","CN304","CN363","CN328"]),
]

# ── Status ────────────────────────────────────────────────────────────────────
STATUS_ORDER = {"已修復完成":0,"故障":1,"維修中":2,"待確認":3}
STATUS_ZH_EN = {
    "故障":      ("故障",     "Fault"),
    "已修復完成": ("已修復完成","Fixed"),
    "維修中":    ("維修中",    "Under Repair"),
    "待確認":    ("待確認",    "Pending"),
}
STATUS_COLOR = {
    "故障":      RGBColor(0xC0,0x00,0x00),
    "已修復完成": RGBColor(0x0E,0x74,0x90),
    "維修中":    RGBColor(0x7B,0x3F,0x00),
    "待確認":    RGBColor(0x1F,0x49,0x7D),
}
STATUS_BG = {
    "故障":      "FFF5F5",
    "已修復完成": "F0FDFC",
    "維修中":    "FFFBF0",
    "待確認":    "F5F7FF",
}
ACTIVE_STATUSES = {"故障","維修中","待確認"}

# ── Colors ────────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A,0x23,0x32)
BLUE  = RGBColor(0x2D,0x4A,0x8A)
WHITE = RGBColor(0xFF,0xFF,0xFF)
TEAL  = RGBColor(0x0E,0x74,0x90)
BROWN = RGBColor(0x92,0x40,0x0E)
GRAY  = RGBColor(0xAA,0xAA,0xAA)

# ── Column widths (twips) ─────────────────────────────────────────────────────
C_NO = 267; C_ST = 889; C_ZH = 2312; C_EN = 2934; C_RP = 1156; C_WT = 1156
COL_WIDTHS = [C_NO, C_ST, C_ZH, C_EN, C_RP, C_WT]
PAGE_W = sum(COL_WIDTHS)  # 8714

# ── Complete translation dictionary ──────────────────────────────────────────
ZH_EN = {
    "(1-1)AC main power 工作指示燈沒亮":                "(1-1) AC main power indicator light not lit",
    "(1-11)VoIP 類比電話無法互撥（電話號碼錯誤":         "(1-11) VoIP analog phones cannot call each other (wrong phone numbers)",
    "(1-11)語音伺服器：無號碼表及設備MAC":               "(1-11) Voice server: no number table or device MAC",
    "(1-11)電話號碼錯誤、無各群trunk":                   "(1-11) Wrong phone numbers; no group trunks configured",
    "(1-12)15米纜線綠色（management)不通、30米黃綠藍（管理）不通": "(1-12) 15m green cable (management) not passing; 30m yellow/green/blue management not passing",
    "(1-12)AP的IP設定未檢查":                            "(1-12) AP IP configuration not verified",
    "(1-12)Mgr筆電無法匯入計劃（顯示不支援此格式類型）": "(1-12) MGR laptop cannot import plan (unsupported format type)",
    "(1-12)ping不到無線基地台(CN-AP)":                   "(1-12) Cannot ping wireless AP (CN-AP)",
    "(1-14)&(1-17)DK switch 管理訊號無法連回車廂，光纖燈號沒亮": "(1-14)&(1-17) DK switch management signal cannot connect back to carriage; fiber link light not lit",
    "(1-14)DK無線電紅色管理埠損壞":                      "(1-14) DK radio red management port damaged",
    "(1-15)48port switch 固定位置過深":                   "(1-15) 48-port switch mounting position too deep",
    "(1-15)Cp switch ping 得到但ssh進不去":               "(1-15) CP switch pingable but SSH inaccessible",
    "(1-16)MGR, OPR筆電跑很慢（建議重灌）":              "(1-16) MGR/OPR laptops running very slow (recommend reinstall)",
    "(1-16)Opr登不進最高權限":                           "(1-16) OPR cannot log in with admin privileges",
    "(1-16)Service desk帳號無法登入":                    "(1-16) Service desk account cannot log in",
    "(1-17)15米饋線紅色（DATA)損壞":                     "(1-17) 15m feeder cable red (DATA) damaged",
    "(1-17)DK無線電網頁常當機":                          "(1-17) DK radio web interface frequently crashes",
    "(1-17)DK紅無線電機 DATA埠口無法進入無線電機":        "(1-17) DK red radio DATA port cannot access radio unit",
    "(1-17)使用dwbk無線電機網頁會一直當機":               "(1-17) DWBK radio web interface keeps crashing",
    "(1-17)區sw 3port 損壞（改置第7port)":               "(1-17) Zone switch port 3 damaged (relocated to port 7)",
    "(1-17)天線桿第二節表面有鋼纜摩擦過痕跡":             "(1-17) Antenna mast 2nd section shows steel cable friction marks",
    "(1-19)MissionOPs 監管時，GPS無資訊":                "(1-19) MissionOps monitoring: no GPS information",
    "(1-19)MissionOPs監管時無GPS資訊":                   "(1-19) MissionOps monitoring: no GPS information",
    "(1-19)MissionOPs監管時，無法取得GPS資訊":            "(1-19) MissionOps monitoring: unable to obtain GPS information",
    "(1-2)天線桿s鉤變形":                                "(1-2) Antenna mast S-hook deformed",
    "(1-5)RR時區未同步(因ap機架損壞無gps）":              "(1-5) RR time zone not synced (no GPS due to AP rack damage)",
    "(1-5)RR無法用ssh進去看組態、ping得到進不去":         "(1-5) RR SSH inaccessible for config; pingable but cannot connect",
    "(1-6)&(1-17)Red switch第19 port的網路線異常（不會亮燈）": "(1-6)&(1-17) Red switch port 19 cable abnormal (no link light)",
    "(1-7)POE供電集線器：port1~4供電異常；POE fail, Fan1.2 fail": "(1-7) PoE hub: ports 1–4 power abnormal; POE fail, Fan 1/2 fail",
    "(1-7)POE供電集線器：燈號未亮，未供電（只供綠色）":   "(1-7) PoE hub: indicator not lit, no power delivery (green only)",
    "(1-9)DNS/MOPs未指向NMS":                            "(1-9) DNS/MOPs not pointing to NMS",
    "(1-9)DNS時間未指向NMS":                             "(1-9) DNS time not synchronized to NMS",
    "(1-9)DNS未指向NMS":                                 "(1-9) DNS not pointing to NMS",
    "(1-9)OPR筆電時間未指向NMS":                         "(1-9) OPR laptop time not synchronized to NMS",
    "(1-9)伺服器/DNS/Mops時間未指向NMS":                 "(1-9) Server/DNS/MOPs time not synchronized to NMS",
    "(1-9)伺服器本機時間未同步NMS":                       "(1-9) Server local time not synchronized to NMS",
    "(?)電燈開關按鍵損壞(缺失表未登載)":                  "(?) Light switch button damaged (not in deficiency list)",
    "15米Management無法連線(藍色)":                      "15m management cable (blue) cannot connect",
    "15米纜線(綠)故障":                                   "15m cable (green) faulty",
    "15米纜線的藍綠損壞":                                 "15m cable blue/green conductors damaged",
    "15米纜線紅色DATA不通":                               "15m cable red (DATA) not passing signal",
    "15米纜線藍MGMT無法使用":                             "15m cable blue MGMT unusable",
    "15米纜線藍、黃 management 訊號沒過":                 "15m cable blue/yellow management signal not passing",
    "15米纜線藍控線頭脫落訊號不過":                       "15m cable blue management connector detached; signal not passing",
    "15米纜線黃MGMT無法使用":                             "15m cable yellow MGMT unusable",
    "15米電纜線綠色控制故障、紅色控制故障、延長線無線盤":  "15m cable green/red control faulty; extension cord reel missing",
    "15米饋線藍色radio端data接頭損壞":                    "15m feeder cable blue radio-end data connector damaged",
    "30米紅綠MGMT接頭損壞":                               "30m cable red/green MGMT connectors damaged",
    "5支電話未建立、voip（10E376013564、505C885C6DFB）無法開機": "5 phones not provisioned; VoIP units cannot power on",
    "Agent無法載入":                                      "Agent cannot load",
    "Ap ping 不到":                                       "AP not pingable",
    "Ap打不開頁面":                                       "AP web page cannot be opened",
    "BLK SW無console接口":                               "BLK switch has no console port",
    "CN 黃色無線電的Poe埠無法鎖緊(紅點會露出）":          "CN yellow radio PoE port cannot lock (red dot exposed)",
    "CP SW在MOPS監控無法顯示":                            "CP switch not displayed in MOPs monitoring",
    "CP箱UPS故障#301":                                   "CP box UPS failure #301",
    "CP箱蓋子卡楯損壞(front)":                           "CP box lid latch broken (front)",
    "Case缺2條短網路線、DWBK 綠無線電缺Poe防塵蓋\n(網路線已補）": "Case missing 2 short network cables; DWBK green radio missing PoE dust cap (cables replenished)",
    "DK交換器需重灌組態":                                 "DK switch configuration needs to be reloaded",
    "DK紅色無線電機gps孔少防塵套":                        "DK red radio GPS port missing dust cap",
    "DK綠色無線電機頻繁當機":                             "DK green radio unit frequently crashes",
    "DK綠色無線電機，管理IP網頁無法開啟\n6/26 Brandon 查為15米cable壞掉": "DK green radio management web page cannot open (6/26 Brandon: 15m cable faulty)",
    "DK黃色無線電機無綠色指示燈":                         "DK yellow radio unit green indicator light missing",
    "DNS/MOPS虛擬機時間未指向NMS":                        "DNS/MOPs VM time not synchronized to NMS",
    "DNS/MOPS虛擬機時間未指向nms":                        "DNS/MOPs VM time not synchronized to NMS",
    "DNS/MOPs時區未指向Host":                             "DNS/MOPs time zone not pointing to Host",
    "DNS/MOPs未指向NMS":                                  "DNS/MOPs not pointing to NMS",
    "DNS/MOPs未指向NMS（待確認）":                        "DNS/MOPs not pointing to NMS (pending confirmation)",
    "DNS時間未指向NMS":                                   "DNS time not synchronized to NMS",
    "DNS未指向NMS":                                       "DNS not pointing to NMS",
    "DNS未指向NMS（待確認）":                             "DNS not pointing to NMS (pending confirmation)",
    "DNS未指向NMS；未完成IP設定":                         "DNS not pointing to NMS; IP configuration incomplete",
    "DNS虛擬機伺服器時間未指向nms":                       "DNS VM server time not synchronized to NMS",
    "DNS虛擬機時間未指向NMS":                             "DNS VM time not synchronized to NMS",
    "DNS虛擬機未指向NMS":                                 "DNS VM not pointing to NMS",
    "DWBK光跳線A1損壞":                                   "DWBK fiber patch cable A1 damaged",
    "DWBK四台無線電機連線GUI時，頁面斷斷續續":             "DWBK four radios: GUI connection intermittent/unstable",
    "DWBK的無線電機接在DWBK上時，所有無線電機都無法用管理埠口登入。": "When DWBK radios are connected to DWBK, all radios cannot be accessed via management port",
    "DWBK的紅色無線電機，直接接在車廂上時，管理埠口無法登入。使用26.2登入後修改內容，無法儲存。儲存後登出再登入又還原成原狀。已嘗試重置無線電機。": "DWBK red radio (direct carriage connection): management port login fails; changes cannot be saved and revert after re-login. Reset attempted.",
    "ECU功能正常，但有異音":                               "ECU functioning normally but producing abnormal noise",
    "ECU加熱功能無作用":                                  "ECU heating function not working",
    "ECU故障":                                            "ECU failure",
    "ECU無法製冷(H8)":                                    "ECU cooling function not working (H8)",
    "ECU熱功能無作用":                                    "ECU heating function not working",
    "Exchange服務若沒有Root伺服器情況下無法使用":           "Exchange service cannot function without Root server",
    "MGR筆電mission ops無法讀取":                         "MGR laptop: MissionOps cannot read data",
    "MGR筆電時間未指向NMS":                               "MGR laptop time not synchronized to NMS",
    "MGR筆電未安裝putty":                                 "MGR laptop: PuTTY not installed",
    "MGR缺tftp軟體":                                      "MGR missing TFTP software",
    "MOPS虛擬機時間未指向NMS":                            "MOPs VM time not synchronized to NMS",
    "MOPs時間未指向NMS":                                   "MOPs time not synchronized to NMS",
    "MOPs未指向NMS":                                       "MOPs not pointing to NMS",
    "MOPs未指向NMS；未完成IP設定":                         "MOPs not pointing to NMS; IP configuration incomplete",
    "Mgr筆電、伺服器時間未指向nms":                        "MGR laptop and server time not synchronized to NMS",
    "MissionOps GPS定位無法抓取":                          "MissionOps unable to acquire GPS location",
    "Missionops異常緩慢":                                  "MissionOps running abnormally slow",
    "Mops 虛擬機時間未指向NMS":                            "MOPs VM time not synchronized to NMS",
    "NAS未完成RAID5設定":                                  "NAS RAID5 configuration not completed",
    "OPR筆電未指向NMS":                                    "OPR laptop not pointing to NMS",
    "OPR電腦outlook網頁無法\n登入":                        "OPR laptop: Outlook web cannot log in",
    "Opr筆電未指向NMS":                                    "OPR laptop not pointing to NMS",
    "Outlook應用程式無法連線":                             "Outlook application cannot connect",
    "Outlook網頁可進入頁面，無法登入":                     "Outlook web page accessible but cannot log in",
    "POE Fan1.2 fail":                                    "PoE Fan 1/2 failure",
    "POE fail, Fan1.2亮紅燈；POE in use 1,2,3,4未亮綠燈": "PoE fail; Fan 1/2 red light; PoE in-use ports 1–4 not lit green",
    "POE fail亮燈，Fan1.2亮燈；port 3無法供電":            "PoE fail light on, Fan 1/2 light on; port 3 cannot supply power",
    "POE供電集線器功能失效":                               "PoE power hub function failed",
    "POE供電集線器：POE Fan1.2 fail; 無接線（在壁櫃）":    "PoE hub: Fan 1/2 fail; no cabling (stored in cabinet)",
    "POE供電集線器：POE fail, port3.4不過電，風扇損壞，燈號異常": "PoE hub: POE fail, ports 3/4 no power, fan damaged, indicator abnormal",
    "POE供電集線器：帳號無法登入":                         "PoE hub: account cannot log in",
    "POE供電集線器：風扇和POE fail故障(要拿備料更換）":     "PoE hub: fan and POE fail (spare parts needed for replacement)",
    "POE功能異常: POE fail, fan1.2":                      "PoE function abnormal: POE fail, Fan 1/2",
    "POE登入帳號錯誤":                                    "PoE login credentials incorrect",
    "POE第三port燈號不會亮，但ping的到":                   "PoE port 3 indicator not lit but pingable",
    "Ping不到無線AP，無法登入":                            "Wireless AP not pingable and cannot log in",
    "Poe fan 亮紅燈":                                     "PoE fan indicator red",
    "RF-9850紅（車頂）管理埠(J31)損壞，無線電機無法進入":  "RF-9850 red (rooftop) management port (J31) damaged; radio unit inaccessible",
    "RJ45端 黃色管理接頭故障（外觀正常）":                 "RJ45 end yellow management connector faulty (appears normal externally)",
    "Trunk未完成，用戶未建 7/1檢查":                       "Trunk not configured, users not created; check on 7/1",
    "UPS跳故障碼":                                        "UPS triggering fault code",
    "VoIP(PNA30-0375-121)話筒線無法卡入":                 "VoIP (PNA30-0375-121) handset cord cannot lock in",
    "VoIP可撥打類比電話，但無法回撥":                      "VoIP can call analog phone but cannot receive return calls",
    "VoIP無法與類比電話互撥":                              "VoIP cannot make/receive calls with analog phones",
    "Zabbix: ping不到":                                   "Zabbix: host not pingable",
    "mission ops GPS定位無法抓取":                         "MissionOps unable to acquire GPS location",
    "opr筆電畫面黑屏":                                    "OPR laptop screen blank",
    "ping不到無線基地台AP":                                "Wireless AP not pingable",
    "中繼交換器vlan105 顯示down down":                     "Relay switch VLAN 105 shows down/down",
    "伺服器最高權限帳號為L3Harris":                        "Server admin account is L3Harris (non-standard)",
    "伺服器本機/DNS/MOPs未指向NMS（待確認）":               "Server local/DNS/MOPs not pointing to NMS (pending confirmation)",
    "伺服器本機時間未指向NMS":                             "Server local time not synchronized to NMS",
    "伺服器本機未指向NMS":                                 "Server local not pointing to NMS",
    "使用dwbk無線電機網頁會一直當機":                      "Using DWBK radio web interface keeps crashing",
    "冷氣上午打不開下午又恢復正常":                        "Air conditioning cannot start in morning; recovers in afternoon",
    "冷氣無法開啟（待確認）":                              "Air conditioning cannot turn on (pending confirmation)",
    "因訊號入口面板J4 埠口燒毀導致紅色路由器G1/0/1/顯示down down": "Signal entry panel J4 port burned out; red router G1/0/1 shows down/down",
    "工作站筆電無標籤（MGR, OPR）":                        "Workstation laptops unlabeled (MGR, OPR)",
    "時間伺服器GPS為2D":                                   "Time server GPS fix is 2D (insufficient)",
    "時間伺服器：\nGPS fix: fix not acquired\nSatellites 沒大於4": "Time server: GPS fix not acquired; satellites fewer than 4",
    "本機ping不到moxa (192.168.255.235)網頁進不去":        "Local host cannot ping MOXA (192.168.255.235); web page inaccessible",
    "無法撥打類比電話":                                    "Cannot make analog phone calls",
    "無線基地台AP未完成IP設定":                            "Wireless AP IP configuration incomplete",
    "無線基地台AP未完成IP設定，網頁進不去":                 "Wireless AP IP configuration incomplete; web page inaccessible",
    "無線電端 紅色管理接頭故障（pin腳彎曲）":               "Radio end red management connector faulty (bent pins)",
    "發電機保養燈":                                        "Generator maintenance indicator light on",
    "發電機啟始功能正常，但面板會顯示系統異常燈":            "Generator starts normally but panel shows system error light",
    "發電機面板側右下滑軌卡楯不回彈":                      "Generator panel right lower slide rail latch does not spring back",
    "發電機面板側滑軌安全孔異常":                          "Generator panel slide rail safety hole abnormal",
    "筆電網路線少一條":                                    "Laptop network cable missing",
    "紅色交換器第十條網路線損壞":                          "Red switch 10th network cable damaged",
    "紅色路由器FXS第一port異常，voip跟類比電話都無法互播": "Red router FXS port 1 abnormal; VoIP and analog phones cannot communicate",
    "網路線有一條是壞掉的，所以LACP只啟始一個埠口":         "One network cable faulty; LACP only brings up one port",
    "缺延長線及延長線盤":                                  "Extension cord and reel missing",
    "缺電纜連接頭":                                        "Missing cable connector",
    "藍色radio管理埠失效\n6/26 Brandon 查為15米cable壞掉":  "Blue radio management port failed (6/26 Brandon: 15m cable faulty)",
    "藍色基頻纜線POE失效（243）":                          "Blue baseband cable PoE failure (243)",
    "視訊電話無法成功註冊":                                "Video phone cannot register successfully",
    "語音伺服器未建立用戶":                                "Voice server: users not created",
    "語音伺服器無Baseline設定，需重載還原檔":               "Voice server: no baseline configuration; restore file reload required",
    "語音伺服器無建立電話":                                "Voice server: phones not provisioned",
    "語音伺服器：Port1亮紅燈；無用戶且無法撥打類比電話":    "Voice server: Port 1 red light; no users provisioned, cannot make analog calls",
    "語音伺服器：內建無電話用戶；撥打類比電話異常":          "Voice server: no phone users provisioned; analog call abnormal",
    "車長側的發電機右側滑軌固定扣無法正常解鎖":             "Commander-side generator right slide rail latch cannot unlock normally",
    "車頂綠radio缺1防塵蓋、30米缺板手":                    "Rooftop green radio missing 1 dust cap; 30m cable missing wrench",
    "電源分配箱AC main power指示燈未亮":                   "Power distribution box AC main power indicator not lit",
    "電話攜行箱：號碼清單正常但電話無法登入":               "Phone carry case: number list normal but phone cannot log in",
    "電話清單不完整、類比電話號碼有錯":                     "Phone list incomplete; analog phone numbers incorrect",
    "需重載Sigma還原檔":                                   "Sigma restore file reload required",
    "類比電話的號碼不是該車廂的號碼":                       "Analog phone number does not match carriage's assigned number",
    "類比電話的號碼不是該車廂的電話號碼":                   "Analog phone number does not match carriage's assigned phone number",
    # ── Additional translations ───────────────────────────────────────────────
    "(1-11)16支電話可撥打，其中5支加入會議室":               "(1-11) 16 phones operational; 5 added to conference room",
    "(1-11)16支電話可正常撥號，並將其中五支加入會議":         "(1-11) 16 phones dial normally; 5 added to conference",
    "(1-11)16支電話均可撥打，其中5支加入會議室":             "(1-11) All 16 phones operational; 5 added to conference room",
    "(1-11)VOIP可撥打類比電話":                             "(1-11) VoIP can call analog phones",
    "(1-11)可撥打類比電話":                                 "(1-11) Can make analog phone calls",
    "(1-11)無法撥打類比電話":                               "(1-11) Cannot make analog phone calls",
    "(1-11)語音伺服器未建立用戶":                           "(1-11) Voice server: users not created",
    "(1-11)需重載Sigma還原檔":                              "(1-11) Sigma restore file reload required",
    "(1-12)Ping不到無線AP，無法登入":                       "(1-12) Wireless AP not pingable and cannot log in",
    "(1-12)無線基地台AP未完成IP設定":                       "(1-12) Wireless AP IP configuration incomplete",
    "(1-14)紅藍Dk無線電導波管垂直接頭變形":                 "(1-14) DK red/blue radio waveguide vertical connector deformed",
    "(1-16)Opr筆電的MOPS單機版無法進入":                    "(1-16) OPR laptop: MOPs standalone version inaccessible",
    "(1-17)DK藍色無線電機接DWBK網頁當機頻繁":               "(1-17) DK blue radio connected to DWBK: web interface frequently crashes",
    "(1-17)天線杆S扣鉤變形":                                "(1-17) Antenna mast S-hook deformed",
    "(1-17)天線桿第四、五節有劃痕無異音":                   "(1-17) Antenna mast sections 4/5 show scratches; no abnormal noise",
    "(1-17)紅區sw 3port 損壞（改置第7port)":                "(1-17) Red zone switch port 3 damaged (relocated to port 7)",
    "(1-19)Agent要用CSV匯出再匯入才能載":                   "(1-19) Agent requires CSV export/import to load plan",
    "(1-19)DK藍色無線電機連接DWBK網頁當機頻繁":             "(1-19) DK blue radio connected to DWBK: web interface frequently crashes",
    "(1-19)plan要從新匯出csv 檔再匯入missionops  agent才有辦法載入": "(1-19) Plan must be re-exported as CSV and re-imported to MissionOps Agent to load",
    "(1-2)冷氣無法開啟（上午無法開啟，下午可以開啟）":       "(1-2) Air conditioning cannot turn on in morning; works in afternoon",
    "(1-2)冷氣異常，切換風量最大及Auto時會停止運作":         "(1-2) Air conditioning abnormal; stops when switched to max fan or Auto mode",
    "(1-6)&(1-12)因訊號入口面板J4 埠口燒毀導致紅色路由器G1/0/17顯示down down": "(1-6)&(1-12) Signal entry panel J4 port burned out; red router G1/0/17 shows down/down",
    "(1-7)POE fail, Fan1.2亮紅燈；POE in use 1,2,3,4未亮綠燈": "(1-7) PoE fail; Fan 1/2 red light; PoE in-use ports 1–4 not lit green",
    "(1-8)POE功能異常: POE fail, fan1.2":                   "(1-8) PoE function abnormal: POE fail, Fan 1/2",
    "(1-9)DNS/MOPs時區未指向Host":                          "(1-9) DNS/MOPs time zone not pointing to Host",
    "(1-9)DNS未指向NMS（待確認）":                          "(1-9) DNS not pointing to NMS (pending confirmation)",
    "(1-9)MOPS服務有啟動，但網頁無法進入（鳥會一直飛）":    "(1-9) MOPs service started but web page inaccessible (loading animation loops)",
    "30米電纜線無線電端mgt接頭防塵蓋遺失":                  "30m cable radio-end MGT connector dust cap missing",
    "4台DK無線電機連接DBWK網頁當機頻繁":                    "4 DK radios connected to DWBK: web interface frequently crashes",
    "DK藍色無線電機頁面，按SAVE後會跳出Rebooting指示，但不會重開機。": "DK blue radio web page: pressing SAVE shows Rebooting prompt but device does not reboot",
    "MOPS監管時，無GPS資訊":                                "MOPs monitoring: no GPS information",
    "UPS故障 #304":                                         "UPS failure #304",
    "UPS故障碼#304":                                        "UPS fault code #304",
    "pcase voip缺1背板、缺網路線3":                         "Pcase: VoIP missing 1 backplate; 3 network cables missing",
    "pcase少2條網路線":                                     "Pcase: 2 network cables missing",
    "伺服器無法進入，iLo也無法":                            "Server inaccessible; iLO also inaccessible",
    "冷氣上午打不開下午又恢復正常(時好時壞)［用發電機送電後正常］": "Air conditioning cannot start in morning; recovers in afternoon (intermittent) [normal after generator power]",
    "天線桿第二三四節有畫痕":                               "Antenna mast sections 2/3/4 show scratches",
    "第三、四、五節有劃痕無異音":                           "Sections 3/4/5 show scratches; no abnormal noise",
    "紅區路由器狀態燈亮紅燈":                               "Red zone router status indicator red",
    "網路線損壞*1，連接筆電顯示unplugged，連接話機可開機無法註冊（放在車廂左側抽屜）": "Network cable damaged ×1; shows unplugged on laptop; phone powers on but cannot register (stored in left drawer)",
    "車頂置物箱MAP兩支都缺天線網的固定卡榫、一支螺絲無法鎖緊。": "Rooftop storage box MAP: both units missing antenna bracket retaining clips; one screw cannot tighten",
    "電燈開關按鍵損壞":                                     "Light switch button damaged",
    "（1-11）VoIP可撥打類比電話，但無法回撥":               "(1-11) VoIP can call analog phone but cannot receive return calls",
    "（1-11）VoIP無法與類比電話互撥":                       "(1-11) VoIP cannot make/receive calls with analog phones",
    "（1-12）15米纜線黃MGMT無法使用":                       "(1-12) 15m cable yellow MGMT unusable",
    "（1-12）30米紅綠MGMT接頭損壞":                         "(1-12) 30m cable red/green MGMT connectors damaged",
    "（1-12）無線基地台AP未完成IP設定":                     "(1-12) Wireless AP IP configuration incomplete",
    "（1-12）無線基地台AP未完成IP設定，網頁進不去":          "(1-12) Wireless AP IP configuration incomplete; web page inaccessible",
    "（1-12）藍色基頻纜線POE失效（243）":                   "(1-12) Blue baseband cable PoE failure (243)",
    "（1-12）黃色饋線接頭蓋子損壞（Data端）":               "(1-12) Yellow feeder connector cover damaged (data end)",
    "（1-14）DWBK SWwitch光跳線B2損壞":                     "(1-14) DWBK switch fiber patch cable B2 damaged",
    "（1-15）CP箱蓋子卡楯損壞(front)":                      "(1-15) CP box lid latch broken (front)",
    "（1-16）OPR筆電未指向NMS":                             "(1-16) OPR laptop not pointing to NMS",
    "（1-17）15米綠色饋線損壞（資訊）":                     "(1-17) 15m green feeder cable damaged (data)",
    "（1-17）15米纜線藍MGMT無法使用":                       "(1-17) 15m cable blue MGMT unusable",
    "（1-17）4台DK無線電機連接DBWK網頁當機頻繁":             "(1-17) 4 DK radios connected to DWBK: web interface frequently crashes",
    "（1-17）DK綠色無線機網頁常當機（接車廂時）當機後觀察通聯及OSPF都正常": "(1-17) DK green radio web frequently crashes when connected to carriage; comms and OSPF normal after crash",
    "（1-17）DK藍色無線電機連接DWBK網頁當機頻繁":           "(1-17) DK blue radio connected to DWBK: web interface frequently crashes",
    "（1-17）DK黃色無線電機頻繁當機":                       "(1-17) DK yellow radio unit frequently crashes",
    "（1-17）天線杆S扣鉤變形（道路側）":                    "(1-17) Antenna mast S-hook deformed (road side)",
    "（1-17）天線桿第三節有劃痕":                           "(1-17) Antenna mast section 3 shows scratches",
    "（1-17）天線桿第二、三、四節有劃痕":                   "(1-17) Antenna mast sections 2/3/4 show scratches",
    "（1-17）天線桿第二、四節有劃痕，升降有異音":            "(1-17) Antenna mast sections 2/4 show scratches; abnormal noise during elevation",
    "（1-17）藍色無線電機接DWBK廂網頁會時常當機":           "(1-17) Blue radio connected to DWBK carriage: web interface frequently crashes",
    "（1-19）Agent要用CSV匯出再匯入才能載入":               "(1-19) Agent requires CSV export/import to load plan",
    "（1-19）GPS天線異常時好時壞無法正常鎖定衛星（廠商檢測完確認異常-目前無備品可以做更換）": "(1-19) GPS antenna intermittent fault; cannot lock satellites (vendor confirmed defective — no spare available)",
    "（1-19）MissionOps 單機版無法使用":                    "(1-19) MissionOps standalone version cannot be used",
    "（1-19）Missionops監控載台位置錯誤（顯示為0）":        "(1-19) MissionOps monitoring shows incorrect platform position (displays as 0)",
    "（1-1）AC電源燈不亮":                                  "(1-1) AC power indicator light not lit",
    "（1-2）ECU功能正常，但有異音":                         "(1-2) ECU functioning normally but producing abnormal noise",
    "（1-2）ECU加熱功能無作用":                             "(1-2) ECU heating function not working",
    "（1-4）時間伺服器GPS為2D":                             "(1-4) Time server GPS fix is 2D (insufficient)",
    "（1-4）時間伺服器：\nGPS fix: fix not acquired\nSatellites 沒大於4": "(1-4) Time server: GPS fix not acquired; satellites fewer than 4",
    "（1-8）POE Fan1.2 fail":                               "(1-8) PoE Fan 1/2 failure",
    "（1-8）POE fail亮燈，Fan1.2亮燈；port 3無法供電":      "(1-8) PoE fail light on, Fan 1/2 light on; port 3 cannot supply power",
    "（1-8）POE供電集線器：POE Fan1.2 fail; 無接線（在壁櫃）": "(1-8) PoE hub: Fan 1/2 fail; no cabling (stored in cabinet)",
    "（1-8）POE損壞":                                       "(1-8) PoE unit damaged",
    "（1-8）POE第三port燈號不會亮，但ping的到":              "(1-8) PoE port 3 indicator not lit but pingable",
    "（1-9）DNS/MOPs時區未指向Host":                        "(1-9) DNS/MOPs time zone not pointing to Host",
    "（1-9）DNS/MOPs未指向NMS（待確認）":                   "(1-9) DNS/MOPs not pointing to NMS (pending confirmation)",
    "（1-9）DNS未指向NMS；未完成IP設定":                    "(1-9) DNS not pointing to NMS; IP configuration incomplete",
    "（1-9）MOPs未指向NMS；未完成IP設定":                   "(1-9) MOPs not pointing to NMS; IP configuration incomplete",
    "（1-9）伺服器本機未指向NMS":                           "(1-9) Server local not pointing to NMS",
    "（清點）15米纜線藍控線頭脫落訊號不過":                  "(Inventory) 15m cable blue management connector detached; signal not passing",
    "（清點）交管手電筒紅色套筒損壞":                        "(Inventory) Traffic control flashlight red sleeve damaged",
    "（清點）缺電源線盤":                                   "(Inventory) Power cord reel missing",
    "（起始）車頂置物箱導波管卡榫損壞*1":                    "(Startup) Rooftop storage box waveguide retaining clip damaged ×1",
    "（1-11）Sigma組態電話未全建立":                        "(1-11) Sigma configuration: phones not fully provisioned",
    "（1-12）&（1-17）15米電纜線藍色管理損壞":              "(1-12)&(1-17) 15m cable blue management conductor damaged",
    "（1-12）30米電纜線紅色管理損壞":                       "(1-12) 30m cable red management conductor damaged",
    "（1-12）Radio (192.168.255.240)ping不到":             "(1-12) Radio (192.168.255.240) not pingable",
    "（1-12）未安裝AP（因ap 機架損壞）,導致Red SW port24 down down": "(1-12) AP not installed (AP rack damaged); Red SW port 24 down/down",
    "（1-16）Mgr、opr未指向NMS":                           "(1-16) MGR/OPR not pointing to NMS",
    "（1-17）Tvam升降有異音":                              "(1-17) TVAM elevation mechanism producing abnormal noise",
    "（1-17）天線桿S鉤環變形":                             "(1-17) Antenna mast S-hook ring deformed",
    "（1-17）天線桿有異音（升降正常，需持續觀察）":          "(1-17) Antenna mast producing abnormal noise (elevation normal; monitor continuously)",
    "（1-19）MissionOPs 監管時，GPS無資訊":                "(1-19) MissionOps monitoring: no GPS information",
    "（1-19）MissionOPs 監管時，無GPS資訊":                "(1-19) MissionOps monitoring: no GPS information",
    "（1-19）MissionOPs監管時無GPS資訊":                   "(1-19) MissionOps monitoring: no GPS information",
    "（1-19）MissionOPs監管時，NTP監控異常":               "(1-19) MissionOps monitoring: NTP monitoring abnormal",
    "（1-19）MissionOPs監管時，無GPS資訊":                 "(1-19) MissionOps monitoring: no GPS information",
    "（1-19）MissionOPs監管時，無GPS資訊；Agent需用csv檔匯入計劃才可載入": "(1-19) MissionOps: no GPS; Agent requires CSV file import to load plan",
    "（1-19）網路監管-MissionOPs監控時無GPS資訊":           "(1-19) Network monitoring: MissionOps shows no GPS information",
    "（1-1）Ac main power 燈沒亮":                         "(1-1) AC main power indicator not lit",
    "（1-1）發電機無法發動（顯示：緊急停車）":              "(1-1) Generator cannot start (shows: emergency stop)",
    "（1-1）突波抑制器未亮燈":                             "(1-1) Surge suppressor indicator not lit",
    "（1-6）未安裝gps、RS G1/0/24顯示down down(因ap 機架損壞無gps)": "(1-6) GPS not installed; RS G1/0/24 shows down/down (AP rack damaged)",
    "（1-8）Poe埠口1顯示燈號異常（運作正常但燈號沒亮）":    "(1-8) PoE port 1 indicator abnormal (functioning but light not lit)",
    "（1-9）DNS/MOPs未指向NMS":                            "(1-9) DNS/MOPs not pointing to NMS",
    "（1-9）DNS/mops未指向NMS":                            "(1-9) DNS/MOPs not pointing to NMS",
    "（1-9）DNS及MOPS已指向NMS，但尚未同步":               "(1-9) DNS and MOPs pointing to NMS but not yet synchronized",
    "(1-19)MOPS監管時，無GPS資訊":                          "(1-19) MOPS monitoring: no GPS information",
    "(清點)無橘色轉接頭":                                    "(Inventory) Orange adapter missing",
    "30米電纜線無防塵蓋":                                    "30m cable dust cap missing",
    "DK_SW組態不完整":                                       "DK switch configuration incomplete",
    "DK無線電紅、藍、黃色的mgt無法登入":                     "DK radio (red/blue/yellow) management port: cannot log in",
    "OPR網卡錯誤":                                           "OPR laptop NIC error",
    "VoIP電話無保護殼":                                      "VoIP phone missing protective case",
    "（1-17）DK無線電紅、藍色GW設定253DK無法連線，設定230車廂無法連線": "(1-17) DK red/blue radios: GW set to 253 cannot connect to DK; set to 230 cannot connect to carriage",
    "（1-19）MOPS服務一直當機":                              "(1-19) MOPS service keeps crashing",
    "（1-19）MissionOPs 監管時，GPS無資訊，agent需用csv載入": "(1-19) MissionOps monitoring: no GPS information; Agent requires CSV import to load plan",
    # ── 2026-07-09 新增（73G CN327）──
    "30米缺板手":                                             "30m kit missing wrench",
    "車頂綠radio缺1防塵蓋":                                   "Roof green radio missing 1 dust cap",
    # ── 2026-07-09 新增（74G/75G 清點）──
    "(1-5)路由器未有登入banner":                              "(1-5) Router login banner not configured",
    "(1-9)MOPS功能無法正常使用":                              "(1-9) MOPS function not working properly",
    "(1-9)DNS伺服器名稱未改(還是DC)":                         "(1-9) DNS server name not changed (still DC)",
    "(1-15)沒有base_line":                                    "(1-15) No base_line",
    "(1-4)NTP的SNMPv3相關設定尚未設定":                       "(1-4) NTP SNMPv3 settings not configured",
    "(1-9)MOPS的vm未開啟w32tm功能":                           "(1-9) MOPS VM w32tm service not enabled",
    "Opr 單機版mops打不開":                                   "OPR standalone MOPS cannot open",
    "DK 黃色無線電機管理埠進不去":                            "DK yellow radio management port inaccessible",
    "DWBK光纖有問題":                                         "DWBK fiber optic issue",
    "Mgr筆電還沒檢查":                                        "MGR laptop not yet checked",
    "MOPs W32 tm未開":                                        "MOPS w32tm service not enabled",
    "OPR,MGR筆電sigma client未安裝":                          "OPR/MGR laptops: Sigma client not installed",
    "DC伺服器名稱未改":                                       "DC server name not changed",
    "(1-15)cp case的sw太後面":                                "(1-15) CP case switch positioned too far back",
    "(1-11)client無法載入qr code檔":                          "(1-11) Client cannot load QR code file",
    "(1-4)NTP的SNMPv3相關設定尚未完成設定":                   "(1-4) NTP SNMPv3 settings not yet completed",
    "(1-11)沒有base_line":                                    "(1-11) No base_line",
    "發電機告警s保養":                                        "Generator S-maintenance alarm",
    "電腦（mgr,opr)都沒有sigma 用戶端":                       "MGR/OPR computers missing Sigma client",
    "Opr筆電無法開啟網頁的mops，單機版正常":                  "OPR laptop cannot open web MOPS; standalone version works",
    "15米纜線黃色沒過電":                                     "15m cable (yellow) no power continuity",
    "(1-12)車頂綠色無線電mgt接頭異常":                        "(1-12) Roof green radio management connector abnormal",
    "(1-11)SIGMA無base_line":                                 "(1-11) SIGMA has no base_line",
    "15米電纜線無線電端data接頭缺蓋子":                       "15m cable radio-end data connector missing cap",
    "(1-9)MOPS功能不正常":                                    "(1-9) MOPS function abnormal",
    "(1-12)DK無線電綠色管理埠ping不到(Layer1沒過)":           "(1-12) DK radio green management port ping fails (Layer 1 fail)",
    "(1-12)DK無線電黃色管理埠ping不到(Layer1有過)":           "(1-12) DK radio yellow management port ping fails (Layer 1 OK)",
    "(1-1)發電機顯示S保警告(ER0101-0076)":                    "(1-1) Generator shows S-maintenance warning (ER0101-0076)",
    "Opr 單機版的mops無network 頁籤、網頁版網址是cn329的，打cn443的網址進不去": "OPR standalone MOPS missing Network tab; web MOPS URL belongs to CN329, CN443's URL inaccessible",
    "電源分配箱過電後有異音":                                 "Power distribution box abnormal noise after power-on",
    "MOPS 載台跳警告不能建立無線電":                          "MOPS platform warning: cannot create radio",
    "MOPS W32未啟動":                                         "MOPS w32tm not started",
    "Cp case ups故障碼#500 RBL4T10037 無法蓄電":              "CP case UPS fault code #500 (RBL4T10037) cannot hold charge",
    "DNS VM命名方式錯誤":                                     "DNS VM naming incorrect",
    "MOPS無NETWORKS 可以建立":                                "MOPS: no NETWORKS available to create",
}

def translate(zh):
    zh_clean = zh.strip()
    if zh_clean in ZH_EN:
        return ZH_EN[zh_clean]
    # Partial match fallback — return original if not found
    return zh_clean

# ── Firestore fetch ───────────────────────────────────────────────────────────
def fetch_all():
    url = f"{BASE_URL}/faultData?key={API_KEY}&pageSize=200"
    req = urllib.request.Request(url, headers={"Accept":"application/json"})
    with urllib.request.urlopen(req, timeout=30) as r:
        data = json.loads(r.read())
    result = {}
    for doc in data.get("documents", []):
        car = doc["name"].split("/")[-1]
        fields = doc.get("fields", {})
        items = []
        for fid, fval in fields.items():
            if fid == "updatedAt": continue
            if "mapValue" not in fval: continue
            m = fval["mapValue"].get("fields", {})
            items.append({
                "id":      fid,
                "status":  m.get("status",  {}).get("stringValue", "待確認"),
                "desc":    m.get("desc",    {}).get("stringValue", ""),
                "person":  m.get("person",  {}).get("stringValue", ""),
                "witness": m.get("witness", {}).get("stringValue", ""),
            })
        if items:
            items.sort(key=lambda x: STATUS_ORDER.get(x["status"], 9))
            result[car] = items
    return result

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_width(cell, w):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(w))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def set_table_fixed(table, widths):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    if tbl.find(qn("w:tblPr")) is None:
        tbl.insert(0, tblPr)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
    tw  = OxmlElement("w:tblW"); tw.set(qn("w:w"), str(sum(widths))); tw.set(qn("w:type"), "dxa"); tblPr.append(tw)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); tblGrid.append(gc)
    tbl.insert(1, tblGrid)

def no_space(para):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing"); sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
    pPr.append(sp)

def add_run(para, text, bold=False, italic=False, color=None, size=8, align=None):
    if align: para.alignment = align
    no_space(para)
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run

# ── Core document builder ─────────────────────────────────────────────────────
def build_doc(fault_data, active_only=False):
    doc = Document()

    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(8)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    # Title
    label_en = "Active Fault Report" if active_only else "Daily Fault Report"
    label_zh = "未修復故障報告" if active_only else "故障日報"
    p = doc.add_paragraph(); no_space(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"DT&E {label_zh} {label_en}"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph(); no_space(p2); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"日期 Date：{TODAY}"); r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x44,0x55,0x66)

    if active_only:
        p3 = doc.add_paragraph(); no_space(p3); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("（僅含故障 / 維修中 / 待確認　Active faults only: Fault / Under Repair / Pending）")
        r3.italic = True; r3.font.size = Pt(8.5); r3.font.color.rgb = RGBColor(0x66,0x66,0x66)

    # Summary
    def count(st):
        return sum(1 for car in fault_data for f in fault_data[car] if f["status"]==st and (not active_only or st in ACTIVE_STATUSES))

    total_cars = sum(1 for car in [c for r in ROTATIONS for c in r[1]]
                     if car in fault_data and any(f["status"] in ACTIVE_STATUSES for f in fault_data[car]))  \
                 if active_only else \
                 sum(1 for car in [c for r in ROTATIONS for c in r[1]] if car in fault_data)

    cnt_fault = count("故障"); cnt_fixed = count("已修復完成")
    cnt_wip   = count("維修中"); cnt_pend  = count("待確認")

    sp = doc.add_paragraph(); no_space(sp)
    sr = sp.add_run("摘要 Summary"); sr.bold = True; sr.font.size = Pt(11); sr.font.color.rgb = NAVY

    TOTAL_CARS = sum(len(r[1]) for r in ROTATIONS)  # 57

    if active_only:
        SUM_COLS = [
            ("總車數\nTotal Cars",          str(TOTAL_CARS), "2C3E50"),
            ("有紀錄車廂\nCars w/ Records", str(total_cars), "1A2332"),
            ("故障\nFault",                 str(cnt_fault),  "C00000"),
            ("維修中\nUnder Repair",        str(cnt_wip),    "7B3F00"),
            ("待確認\nPending",             str(cnt_pend),   "1F497D"),
        ]
        stbl = doc.add_table(rows=2, cols=5)
    else:
        SUM_COLS = [
            ("總車數\nTotal Cars",          str(TOTAL_CARS), "2C3E50"),
            ("有紀錄車廂\nCars w/ Records", str(total_cars), "1A2332"),
            ("故障\nFault",                 str(cnt_fault),  "C00000"),
            ("已修復完成\nFixed",            str(cnt_fixed),  "0E7490"),
            ("維修中\nUnder Repair",        str(cnt_wip),    "7B3F00"),
            ("待確認\nPending",             str(cnt_pend),   "1F497D"),
        ]
        stbl = doc.add_table(rows=2, cols=6)
    set_table_fixed(stbl, [PAGE_W // len(SUM_COLS)] * len(SUM_COLS))
    stbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    cw = PAGE_W // len(SUM_COLS)
    for i,(label,val,bg) in enumerate(SUM_COLS):
        hc = stbl.rows[0].cells[i]; set_cell_bg(hc, bg); set_cell_width(hc, cw)
        hc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hp = hc.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(hp)
        hr = hp.add_run(label); hr.bold = True; hr.font.size = Pt(7.5); hr.font.color.rgb = WHITE

        vc = stbl.rows[1].cells[i]; set_cell_width(vc, cw)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vp = vc.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(vp)
        vr = vp.add_run(val); vr.bold = True; vr.font.size = Pt(18); vr.font.color.rgb = NAVY

    doc.add_paragraph()

    # Rotations
    for rot_label, cars in ROTATIONS:
        # Rotation header
        rt = doc.add_table(rows=1, cols=1); set_table_fixed(rt, [PAGE_W])
        rc = rt.rows[0].cells[0]; set_cell_bg(rc, "1A2332"); set_cell_width(rc, PAGE_W)
        rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.LEFT; no_space(rp)
        rr = rp.add_run(f"DT&E {rot_label} ｜ 車廂數 Cars：{len(cars)}")
        rr.bold = True; rr.font.size = Pt(11); rr.font.color.rgb = WHITE

        for car in cars:
            items = fault_data.get(car, [])
            if active_only:
                items = [f for f in items if f["status"] in ACTIVE_STATUSES]

            # Car header
            ct = doc.add_table(rows=1, cols=2); set_table_fixed(ct, [80, PAGE_W-80])
            bc = ct.rows[0].cells[0]; set_cell_bg(bc, "2D4A8A"); set_cell_width(bc, 80); no_space(bc.paragraphs[0])
            nc = ct.rows[0].cells[1]; set_cell_width(nc, PAGE_W-80)
            # Remove border on name cell
            tcPr = nc._tc.get_or_add_tcPr()
            tcB  = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right"):
                b = OxmlElement(f"w:{side}"); b.set(qn("w:val"),"nil"); b.set(qn("w:sz"),"0"); b.set(qn("w:space"),"0"); b.set(qn("w:color"),"auto"); tcB.append(b)
            tcPr.append(tcB)
            np2 = nc.paragraphs[0]; no_space(np2)
            nr  = np2.add_run(f" {car}"); nr.bold = True; nr.font.size = Pt(11); nr.font.color.rgb = NAVY

            if not items:
                if not active_only:
                    p = doc.add_paragraph(); no_space(p)
                    r = p.add_run("  無故障紀錄 No fault records"); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GRAY
                # In active_only mode, skip cars with no active faults silently
                continue

            tbl = doc.add_table(rows=1+len(items), cols=6)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_fixed(tbl, COL_WIDTHS)

            hdrs = ["#","狀態\nStatus","故障描述 (中文)","Fault Description (English)","修復人員\nRepaired by","見證人\nWitness"]
            for j,h in enumerate(hdrs):
                c = tbl.rows[0].cells[j]; set_cell_bg(c,"2D4A8A"); set_cell_width(c,COL_WIDTHS[j])
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p)
                r = p.add_run(h); r.bold = True; r.font.size = Pt(7.5); r.font.color.rgb = WHITE

            for idx,f in enumerate(items):
                row = tbl.rows[1+idx]
                bg  = STATUS_BG.get(f["status"],"FFFFFF")
                zh_st, en_st = STATUS_ZH_EN.get(f["status"],(f["status"],""))
                sc  = STATUS_COLOR.get(f["status"], NAVY)

                # # col
                c0 = row.cells[0]; set_cell_bg(c0,bg); set_cell_width(c0,C_NO)
                c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p0)
                r0 = p0.add_run(str(idx+1)); r0.font.size = Pt(8); r0.font.color.rgb = RGBColor(0x55,0x55,0x55)

                # Status col
                c1 = row.cells[1]; set_cell_bg(c1,bg); set_cell_width(c1,C_ST)
                c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p1)
                r1a = p1.add_run(zh_st); r1a.bold = True; r1a.font.size = Pt(8); r1a.font.color.rgb = sc
                r1a.add_break()
                r1b = p1.add_run(en_st); r1b.font.size = Pt(7); r1b.font.color.rgb = RGBColor(0x66,0x66,0x66)

                # ZH desc
                c2 = row.cells[2]; set_cell_bg(c2,bg); set_cell_width(c2,C_ZH)
                c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p2 = c2.paragraphs[0]; no_space(p2)
                r2 = p2.add_run(f["desc"]); r2.font.size = Pt(8); r2.font.color.rgb = NAVY

                # EN desc
                c3 = row.cells[3]; set_cell_bg(c3,bg); set_cell_width(c3,C_EN)
                c3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = c3.paragraphs[0]; no_space(p3)
                r3 = p3.add_run(translate(f["desc"])); r3.font.size = Pt(8); r3.font.color.rgb = NAVY

                # Person
                c4 = row.cells[4]; set_cell_bg(c4,bg); set_cell_width(c4,C_RP)
                c4.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p4 = c4.paragraphs[0]; p4.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p4)
                pt  = f["person"].strip()
                if pt:
                    rr4 = p4.add_run(pt); rr4.bold = True; rr4.font.size = Pt(8); rr4.font.color.rgb = TEAL
                else:
                    rr4 = p4.add_run("—"); rr4.font.size = Pt(8); rr4.font.color.rgb = GRAY

                # Witness
                c5 = row.cells[5]; set_cell_bg(c5,bg); set_cell_width(c5,C_WT)
                c5.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p5 = c5.paragraphs[0]; p5.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p5)
                wt  = f["witness"].strip()
                if wt:
                    rr5 = p5.add_run(wt); rr5.bold = True; rr5.font.size = Pt(8); rr5.font.color.rgb = BROWN
                else:
                    rr5 = p5.add_run("—"); rr5.font.size = Pt(8); rr5.font.color.rgb = GRAY

        doc.add_paragraph()

    # Footer
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(fp)
    fr = fp.add_run(f"本報告由系統自動產生 / Auto-generated on {TODAY}")
    fr.italic = True; fr.font.size = Pt(7.5); fr.font.color.rgb = GRAY

    return doc

# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Fetching Firestore data...")
    fault_data = fetch_all()
    print(f"  Got data for {len(fault_data)} cars")

    # Check untranslated
    missing = set()
    for items in fault_data.values():
        for f in items:
            d = f["desc"].strip()
            if d and d not in ZH_EN:
                missing.add(d)
    if missing:
        print(f"  WARNING: {len(missing)} untranslated descriptions (will show original)")

    out1 = f"{SCRATCHPAD}/故障日報_{TODAY}.docx"
    build_doc(fault_data, active_only=False).save(out1)
    print(f"OK: {out1}")
